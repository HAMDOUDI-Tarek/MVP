import re
from docx import Document

def read_blocks(file_path):
    doc = Document(file_path)
    blocks = {}
    current_block = []
    current_code = None
    code_pattern = re.compile(r'^\d+/\d+')

    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "":
            if current_block and current_code:
                if current_code not in blocks:
                    blocks[current_code] = []
                # Skip the first line which contains the code
                blocks[current_code].append("\n".join(current_block[1:]))
                current_block = []
                current_code = None
        elif code_pattern.match(text):
            current_code = code_pattern.match(text).group(0)
            current_block.append(text)
        elif current_code:
            current_block.append(text)

    if current_block and current_code:
        if current_code not in blocks:
            blocks[current_code] = []
        blocks[current_code].append("\n".join(current_block[1:]))

    return blocks

def read_titles_with_code(file_path):
    doc = Document(file_path)
    titles = {}
    code_pattern = re.compile(r'^\d+/\d+')

    for para in doc.paragraphs:
        text = para.text.strip()
        if text and code_pattern.match(text):
            code = code_pattern.match(text).group(0)
            titles[code] = text

    return titles

def create_dictionary(source_path, target_path):
    blocks = read_blocks(source_path)
    titles = read_titles_with_code(target_path)

    result = {}
    for code, title in titles.items():
        if code in blocks:
            result[title] = blocks[code]
        else:
            result[title] = []

    return result

def extract_year(text):
    year_pattern = re.compile(r'20\d{2}')
    lines = text.split("\n")
    for line in lines:
        match = year_pattern.search(line)
        if match:
            return match.group(0)
    return None

def update_target_document(result_dict, target_path):
    doc = Document(target_path)
    code_pattern = re.compile(r'^\d+/\d+')

    for para in doc.paragraphs:
        text = para.text.strip()
        if text and code_pattern.match(text):
            if text in result_dict:
                # Insert #start right after the title

                sorted_blocks = sorted(result_dict[text], key=extract_year, reverse=True)
                sorted_blocks.reverse()  # Invert the list to paste blocks in reverse order
                previous_year = None
                for block in sorted_blocks:
                    year = extract_year(block)

                    if year and year != previous_year:

                        # Insert the year if it has changed
                        #if str(int(year) - 1) != "2019":
                        year_para = doc.add_paragraph(str(int(year) - 1) + "\n")
                        para._element.getparent().insert(para._element.getparent().index(para._element) + 1, year_para._element)

                    new_para = doc.add_paragraph(block)
                    blank_para = doc.add_paragraph("")

                    para._element.getparent().insert(para._element.getparent().index(para._element) + 1, blank_para._element)
                    para._element.getparent().insert(para._element.getparent().index(para._element) + 1, new_para._element)

                    previous_year = year



    doc.save(target_path)

# Example usage:
# Paths to the source and target files
source_file_path = 'c:/users/Expert Info/Desktop/JALIL/Cardio-respiratoire, système unitaire 2ème année _231203_102405.docx'
target_file_path = 'c:/users/Expert Info/Desktop/JALIL/TITRE - TEST.docx'

# Create the dictionary with titles and corresponding blocks of text
result_dict = create_dictionary(source_file_path, target_file_path)

# Update the target document with the blocks of text
update_target_document(result_dict, target_file_path)

print(f"Updated {target_file_path} with blocks of text.")

